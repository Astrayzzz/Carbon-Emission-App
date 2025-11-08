import streamlit as st
import pandas as pd
import numpy as np
import joblib
import time
import matplotlib.pyplot as plt
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt
from sklearn.model_selection import train_test_split

# --- 4. (NEW) ฟังก์ชันสร้างรีพอร์ต (สำหรับ Feature 3) ---
def generate_monthly_report_for_app(y_true, y_pred, kpi_target_kg=10000):
    """
    สร้างรีพอร์ตและกราฟเพื่อแสดงผลบน Streamlit โดยตรง
    จะ return (text_summary, figure_object)
    """
    
    report_lines = []
    today = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    report_lines.append("="*50)
    report_lines.append(f"รายงานสรุปผล Carbon Performance ประจำเดือน")
    report_lines.append(f"สร้างเมื่อ: {today}")
    report_lines.append("="*50)
    
    # --- คำนวณ KPI ---
    total_actual = np.sum(y_true)
    total_predicted = np.sum(y_pred)
    
    # 1. เทียบกับเป้าหมาย (KPI)
    over_target = total_actual - kpi_target_kg
    report_lines.append(f"\n[ผลการดำเนินงานเทียบกับเป้าหมาย (KPI)]")
    report_lines.append(f"  - เป้าหมาย (KPI): {kpi_target_kg:,.0f} Kg")
    report_lines.append(f"  - ปล่อยจริงทั้งหมด: {total_actual:,.0f} Kg")
    
    if over_target > 0:
        report_lines.append(f"  *** สรุป: สูงกว่าเป้า {over_target:,.0f} Kg ***")
    else:
        report_lines.append(f"  *** สรุป: ต่ำกว่าเป้า {abs(over_target):,.0f} Kg (ยอดเยี่ยม!) ***")
        
    # 2. ความแม่นยำของโมเดล (MAE)
    mae = np.mean(np.abs(y_true - y_pred))
    report_lines.append(f"\n[ความแม่นยำของโมเดลพยากรณ์]")
    report_lines.append(f"  - ค่าเฉลี่ยความผิดพลาด (MAE): {mae:.2f} Kg ต่อการเดินทาง")
    
    final_report_text = "\n".join(report_lines)
    
    # --- สร้างกราฟ (Figure Object) ---
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(y_true, label='Actual Carbon (จริง)', marker='o', alpha=0.7)
    ax.plot(y_pred, label=f'Predicted Carbon (พยากรณ์)', marker='x', linestyle='--')
    ax.axhline(y=np.mean(y_true), color='r', linestyle=':', label=f'Actual Mean ({np.mean(y_true):.0f} kg)')
    ax.set_title("Actual vs. Predicted Carbon (Monthly)")
    ax.set_xlabel("Trip ID / Day")
    ax.set_ylabel("Carbon (kg)")
    ax.legend()
    ax.grid(True)
    
    return final_report_text, fig

# --- 1. ฟังก์ชันโหลดโมเดลและ ENCODERS (หัวใจหลัก) ---
@st.cache_resource
def load_model_and_encoders():
    """
    โหลดทั้งโมเดล XGBoost และ Label Encoders ที่บันทึกไว้
    """
    try:
        model = joblib.load("xgboost_carbon_model.pkl")
        encoders = joblib.load("label_encoders.pkl")
        feature_names = model.get_booster().feature_names
        return model, encoders, feature_names
    except FileNotFoundError:
        st.error("ไม่พบไฟล์ '.pkl'! กรุณารัน 'Project.py' (โค้ดเทรน) ก่อนครับ")
        return None, None, None
    except Exception as e:
        st.error(f"เกิดข้อผิดพลาดในการโหลดโมเดล: {e}")
        return None, None, None

# --- 2. ฟังก์ชันจำลองข้อมูลพยากรณ์ (สำหรับ Feature 1) ---
@st.cache_data
def get_mock_forecast_data():
    """จำลองข้อมูลพยากรณ์ 30 วันสำหรับ Dashboard"""
    days = pd.date_range(start="today", periods=30)
    base = 1000 + np.sin(np.arange(30) * 0.5) * 200
    noise = np.random.randint(-50, 50, size=30)
    forecast = base + noise # <-- (เช็กว่าบรรทัดนี้อยู่)
    df = pd.DataFrame({'Date': days, 'Forecasted Carbon (kg)': forecast})
    df = df.set_index('Date')
    return df

# --- 3. (ทางเลือก) ดึงข้อมูล Capacity จาก Dataset.csv ---
@st.cache_data
def get_vehicle_capacities(dataset_file="Dataset.csv"):
    """
    อ่าน Dataset.csv เพื่อสร้าง "คู่มือ" ว่ารถแต่ละประเภท (vehicle_type) 
    มีความจุ (vehicle_capacity_kg) เท่าไหร่
    """
    try:
        df = pd.read_csv(dataset_file)
        capacity_map = df.drop_duplicates(subset=['vehicle_type'])
        capacity_map = capacity_map.set_index('vehicle_type')['vehicle_capacity_kg'].to_dict()
        return capacity_map
    except FileNotFoundError:
        st.error(f"ไม่พบไฟล์ 'Dataset.csv' (จำเป็นต้องใช้ไฟล์นี้เพื่อดูความจุรถ)")
        return None
    except Exception as e:
        st.error(f"เกิดข้อผิดพลาด ao 'Dataset.csv': {e}")
        return None

# -----------------------------------------------------------------
# ---               START OF STREAMLIT APP                      ---
# -----------------------------------------------------------------

st.set_page_config(layout="wide")
st.title("Smart Carbon Logistics Dashboard (Robust Version)")
st.markdown("แพลตฟอร์มต้นแบบสำหรับพยากรณ์และจำลองการปล่อยคาร์บอน")

# --- โหลดทรัพยากรหลัก ---
model, encoders, feature_names = load_model_and_encoders()
capacities = get_vehicle_capacities()

if not all([model, encoders, feature_names, capacities]):
    st.warning("แอปไม่สามารถทำงานได้ กรุณาตรวจสอบข้อผิดพลาดด้านบน")
    st.stop()


# --- Feature 1: Forecast Dashboard (ใช้ข้อมูลจำลอง) ---
st.header("การพยากรณ์ค่าคาร์บอนใน 30 วันข้างหน้า")
st.markdown("กราฟนี้แสดงแนวโน้มการปล่อยคาร์บอนทั้งหมดที่ 'พยากรณ์' ล่วงหน้า 30 วัน (ข้อมูลจำลอง)")
forecast_df = get_mock_forecast_data() 
st.line_chart(forecast_df)

st.divider() 

# --- Feature 2: What-if Simulator (ขับเคลื่อนด้วย XGBoost จริง) ---
st.header("คำนวณค่าคาร์บอน")
st.markdown("ทดลองวางแผนการขนส่งของคุณ โดยระบบจะคำนวณด้วยโมเดล XGBoost (Robust)")

col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("กรอกข้อมูลการขนส่ง")
    
    vehicle_options = list(encoders['vehicle_type'].classes_)
    fuel_options = list(encoders['fuel_type'].classes_)
    
    distance = st.number_input("ระยะทาง (Km)", min_value=0.0, value=100.0, step=0.1, format="%.2f")
    load_weight = st.number_input("น้ำหนักบรรทุก (Kg)", min_value=0.0, value=1000.0, step=1.0)
    
    selected_vehicle_name = st.selectbox("ประเภทรถ (Vehicle Type)", options=vehicle_options)
    selected_fuel_name = st.selectbox("ประเภทเชื้อเพลิง (Fuel Type)", options=fuel_options)
    
    if st.button("คำนวณค่าคาร์บอน", type="primary"):
        with st.spinner("กำลังคำนวณผลลัพธ์"):
            
            try:
                vehicle_code = encoders['vehicle_type'].transform([selected_vehicle_name])[0]
                fuel_code = encoders['fuel_type'].transform([selected_fuel_name])[0]
            except ValueError:
                st.error("เกิดข้อผิดพลาด: ค่าที่เลือกไม่อยู่ใน Encoder")
                st.stop()

            capacity = capacities.get(selected_vehicle_name, 0)
            if capacity == 0:
                relative_load = 0.0
            else:
                relative_load = load_weight / capacity
            
            # --- (✅ แก้ไข) บีบอัด Input ให้ตรงกับโมเดล ---
            distance_log = np.log1p(distance)
            load_weight_log = np.log1p(load_weight)
            # --- (✅) ---

            # 3. สร้าง DataFrame
            input_data = {
                'distance_km': [distance_log],        # <-- ใช้ค่า log
                'load_weight_kg': [load_weight_log],  # <-- ใช้ค่า log
                'vehicle_type': [vehicle_code],
                'fuel_type': [fuel_code],
                'relative_load': [relative_load]
            }
            
            try:
                input_df = pd.DataFrame(input_data)[feature_names]
            except KeyError:
                st.error(f"เกิดข้อผิดพลาด: ลำดับ Feature (feature_names) ไม่ตรงกัน")
                st.stop()

            # 4. ทำนายด้วยโมเดลจริง!
            predicted_carbon = model.predict(input_df)[0]
            
            time.sleep(1) 

        with col2:
            st.subheader("ผลลัพธ์การจำลอง")
            st.metric(
                label="คาร์บอนที่คาดว่าจะปล่อย (Kg)",
                value=f"{predicted_carbon:.2f} kg",
                delta=f"{(predicted_carbon / distance):.2f} kg/km",
                delta_color="inverse"
            )

st.divider() 

# --- Feature 3: สร้างรายงาน "จริง" (จาก Test Set) ---
st.header("สร้างรายงานประสิทธิภาพโมเดล")
st.warning("ขั้นตอนนี้จะทำการโหลดข้อมูลและประมวลผลใหม่ทั้งหมด อาจใช้เวลาสักครู่")

if st.button("สร้างรายงานสรุปผล", type="primary"):
    
    with st.spinner("กำลังโหลดข้อมูลและประเมวลผลโมเดลจริง..."):
        
        try:
            # --- 1. โหลดข้อมูลดิบ (เหมือน Project.py) ---
            df_report = pd.read_csv("Dataset.csv")
            
            # --- 2. ทำ Feature Engineering (เหมือน Project.py) ---
            df_report['vehicle_type'] = encoders['vehicle_type'].transform(df_report['vehicle_type'])
            df_report['fuel_type'] = encoders['fuel_type'].transform(df_report['fuel_type'])
            df_report['relative_load'] = df_report['load_weight_kg'] / df_report['vehicle_capacity_kg']
            
            # --- (✅ แก้ไข) บีบอัดข้อมูลให้ตรงกับโมเดล ---
            df_report['distance_km'] = np.log1p(df_report['distance_km'])
            df_report['load_weight_kg'] = np.log1p(df_report['load_weight_kg'])
            # --- (✅) ---

            # --- 3. สร้าง X และ y (เหมือน Project.py) ---
            X = df_report[feature_names] 
            y = df_report['carbon_kg']

            # --- 4. แบ่งข้อมูล (เหมือน Project.py เป๊ะๆ) ---
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)
            
            # --- 5. ทำนายด้วยโมเดลจริง (เหมือน Project.py) ---
            y_pred = model.predict(X_test)

            time.sleep(1) 

            # --- 7. เรียกใช้ฟังก์ชันรีพอร์ต (ด้วยข้อมูลจริง!) ---
            kpi_target_kg = np.sum(y_test) * 0.9 
            
            report_text, report_fig = generate_monthly_report_for_app(
                y_test.values,
                y_pred,
                kpi_target_kg=kpi_target_kg
            )
            
            st.success("สร้างรายงานจริงเสร็จสิ้น!")
            st.balloons()
            
            st.subheader("รายงานสรุปผล (Summary Report)")
            st.code(report_text) 

            st.subheader("กราฟเปรียบเทียบ (Visual Report)")
            st.pyplot(report_fig)
            
            st.subheader("ดาวน์โหลดรายงาน")

            # 8.1 ปุ่ม .docx
            try:
                doc = Document()
                doc.add_heading('รายงานสรุปผล', level=1)
                doc.add_paragraph(report_text)
                
                buf_docx = BytesIO()
                doc.save(buf_docx)
                st.download_button(
                    label="ดาวน์โหลดรายงาน (.docx)",
                    data=buf_docx.getvalue(),
                    file_name="model_evaluation_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"เกิดข้อผิดพลาดในการสร้าง .docx: {e}")
            
            # 8.2 ปุ่ม .png
            try:
                buf_img = BytesIO()
                report_fig.savefig(buf_img, format="PNG")
                st.download_button(
                    label="ดาวน์โหลดเฉพาะกราฟ (.png)",
                    data=buf_img.getvalue(),
                    file_name="model_evaluation_visual.png",
                    mime="image/png"
                )
            except Exception as e:
                st.error(f"เกิดข้อผิดพลาดในการสร้าง .png: {e}")

        except Exception as e:
            st.error(f"เกิดข้อผิดพลาดร้ายแรงระหว่างการสร้างรายงานจริง: {e}")